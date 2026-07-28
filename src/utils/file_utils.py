"""
Утилиты для работы с файлами
"""

import json
from pathlib import Path
from typing import Any, List, Optional
import pandas as pd
from docx import Document

from .logger import logger


def ensure_dir(path: Path) -> Path:
    """Создает директорию, если её нет"""
    path.mkdir(parents=True, exist_ok=True)
    return path


def save_json(data: Any, filepath: Path, indent: int = 2) -> None:
    """Сохраняет данные в JSON файл"""
    ensure_dir(filepath.parent)
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=indent)
    logger.info(f"JSON сохранен: {filepath}")


def load_json(filepath: Path) -> Any:
    """Загружает данные из JSON файла"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)


def save_csv(df: pd.DataFrame, filepath: Path) -> None:
    """Сохраняет DataFrame в CSV"""
    ensure_dir(filepath.parent)
    df.to_csv(filepath, index=False, encoding='utf-8-sig')
    logger.info(f"CSV сохранен: {filepath}")


def load_csv(filepath: Path) -> pd.DataFrame:
    """Загружает CSV в DataFrame"""
    return pd.read_csv(filepath, encoding='utf-8-sig')


def get_files_by_extension(directory: Path, extension: str) -> List[Path]:
    """Возвращает список файлов с указанным расширением"""
    if not directory.exists():
        logger.warning(f"Директория не существует: {directory}")
        return []
    return list(directory.glob(f"*{extension}"))


# ===== НОВАЯ ФУНКЦИЯ ДЛЯ ЧТЕНИЯ .DOCX =====

def read_docx(filepath: Path) -> str:
    """
    Читает .docx файл и возвращает текст с сохранением структуры таблиц

    Args:
        filepath: Путь к .docx файлу

    Returns:
        str: Текст с таблицами в читаемом виде
    """
    if not filepath.exists():
        logger.error(f"Файл не найден: {filepath}")
        return ""

    try:
        doc = Document(filepath)
        text_parts = []

        # 1. Читаем абзацы
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # 2. Читаем таблицы (сохраняем структуру)
        for table_idx, table in enumerate(doc.tables, 1):
            if not table.rows:
                continue

            text_parts.append(f"\n{'=' * 60}")
            text_parts.append(f"ТАБЛИЦА {table_idx}")
            text_parts.append('=' * 60)

            # Заголовки таблицы (первая строка)
            headers = []
            for cell in table.rows[0].cells:
                headers.append(cell.text.strip().replace('\n', ' '))

            # Определяем ширину колонок для форматирования
            col_widths = [len(h) for h in headers]

            # Проходим по всем строкам для определения макс. ширины
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().replace('\n', ' ')
                    if i < len(col_widths):
                        col_widths[i] = max(col_widths[i], len(cell_text))

            # Формируем разделитель
            separator = "+" + "+".join(["-" * (w + 2) for w in col_widths]) + "+"

            # Выводим заголовки
            header_line = "|"
            for i, header in enumerate(headers):
                header_line += f" {header:<{col_widths[i]}} |"
            text_parts.append(separator)
            text_parts.append(header_line)
            text_parts.append(separator)

            # Выводим данные
            for row in table.rows[1:]:
                row_line = "|"
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().replace('\n', ' ')
                    if i < len(col_widths):
                        row_line += f" {cell_text:<{col_widths[i]}} |"
                text_parts.append(row_line)

            text_parts.append(separator)
            text_parts.append("")

        return "\n".join(text_parts)

    except Exception as e:
        logger.error(f"Ошибка чтения .docx: {e}")
        return f"ERROR: {str(e)}"