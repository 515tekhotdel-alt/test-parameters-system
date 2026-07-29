"""
Экспорт показателей в Excel и Word
"""

import io
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_excel(params_with_methods, rule_info):
    """
    Экспортирует список показателей с методами в Excel
    """
    data = []

    data.append(["Информация о правиле", ""])
    data.append(["Тип изделия", rule_info.get('product_type', '-')])
    data.append(["Возраст", rule_info.get('age', '-')])
    data.append(["Слой", rule_info.get('layer', '-')])
    data.append(["Конструкция", rule_info.get('construction', '-')])
    data.append(["Продуктов в группе", rule_info.get('product_count', 0)])
    data.append(["Совпадение", f"{rule_info.get('score', 0)}%"])
    data.append([])
    data.append(["№", "Контролируемый показатель", "Метод испытаний"])

    for i, item in enumerate(params_with_methods, 1):
        data.append([
            i,
            item.get("name", ""),
            item.get("method", "")
        ])

    df = pd.DataFrame(data)

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Показатели', index=False, header=False)

    return output.getvalue()


def export_to_word(params_with_methods, rule_info):
    """
    Экспортирует список показателей с методами в Word (таблица)
    """
    doc = Document()

    # Заголовок
    title = doc.add_heading('Показатели для испытаний', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Информация о правиле
    doc.add_heading('Информация о продукции', level=1)
    doc.add_paragraph(f'Тип изделия: {rule_info.get("product_type", "-")}')
    doc.add_paragraph(f'Возраст: {rule_info.get("age", "-")}')
    doc.add_paragraph(f'Слой: {rule_info.get("layer", "-")}')
    doc.add_paragraph(f'Конструкция: {rule_info.get("construction", "-")}')
    doc.add_paragraph(f'Продуктов в группе: {rule_info.get("product_count", 0)}')
    doc.add_paragraph(f'Совпадение: {rule_info.get("score", 0)}%')

    doc.add_paragraph()

    doc.add_heading('Контролируемые показатели', level=1)

    # Создаем таблицу
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'

    # Шапка
    header_cells = table.rows[0].cells
    header_cells[0].text = '№'
    header_cells[1].text = 'Контролируемый показатель'
    header_cells[2].text = 'Метод испытаний'

    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True

    # Данные
    for i, item in enumerate(params_with_methods, 1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = item.get("name", "")
        row.cells[2].text = item.get("method", "")

    # Настройка ширины колонок
    for cell in table.columns[0].cells:
        cell.width = Inches(0.5)
    for cell in table.columns[2].cells:
        cell.width = Inches(2.0)
    for cell in table.columns[1].cells:
        cell.width = Inches(4.5)

    # ===== ВСЁ ПО ЛЕВОМУ КРАЮ =====
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()